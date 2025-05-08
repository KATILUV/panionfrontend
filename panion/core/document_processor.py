"""
Document Processor Module
Extracts text and information from various document formats including PDFs, Word documents, and images.
"""

import os
import json
import logging
import tempfile
import base64
from io import BytesIO
from typing import List, Dict, Any, Optional, Union, Tuple, BinaryIO
from datetime import datetime

import pdfplumber
import docx
from PIL import Image
try:
    from PIL import ImageDraw, ImageFont
except ImportError:
    # Optional imports for image annotation
    pass

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Tools for processing and extracting information from documents."""
    
    def __init__(self):
        """Initialize the document processor."""
        self.documents_dir = "./data/documents"
        self.extracted_dir = "./data/extracted"
        self.output_dir = "./data/processed"
        
        # Create necessary directories
        for directory in [self.documents_dir, self.extracted_dir, self.output_dir]:
            os.makedirs(directory, exist_ok=True)
    
    def process_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text and metadata from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        logger.info(f"Processing PDF: {file_path}")
        
        result = {
            "file_name": os.path.basename(file_path),
            "file_path": file_path,
            "file_type": "pdf",
            "pages": [],
            "metadata": {},
            "total_pages": 0,
            "processed_at": datetime.now().isoformat()
        }
        
        try:
            with pdfplumber.open(file_path) as pdf:
                result["total_pages"] = len(pdf.pages)
                
                # Extract metadata
                if hasattr(pdf, 'metadata') and pdf.metadata:
                    result["metadata"] = {
                        k: str(v) for k, v in pdf.metadata.items()
                    }
                
                # Process each page
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text() or ""
                    
                    # Extract tables if available
                    tables = []
                    try:
                        page_tables = page.extract_tables()
                        if page_tables:
                            for j, table in enumerate(page_tables):
                                if table:
                                    table_data = []
                                    for row in table:
                                        if row:
                                            table_data.append([cell or "" for cell in row])
                                    tables.append(table_data)
                    except Exception as e:
                        logger.warning(f"Error extracting tables from page {i+1}: {str(e)}")
                    
                    # Extract images if available
                    images = []
                    try:
                        page_images = page.images
                        if page_images:
                            for j, img in enumerate(page_images):
                                image_data = {
                                    "index": j,
                                    "bbox": img["bbox"]
                                }
                                images.append(image_data)
                    except Exception as e:
                        logger.warning(f"Error extracting images from page {i+1}: {str(e)}")
                    
                    result["pages"].append({
                        "page_number": i + 1,
                        "text": page_text,
                        "tables": tables,
                        "images": images
                    })
                    
                    # Save the extracted text to a file
                    text_filename = f"{os.path.splitext(os.path.basename(file_path))[0]}_page_{i+1}.txt"
                    text_path = os.path.join(self.extracted_dir, text_filename)
                    with open(text_path, 'w', encoding='utf-8') as f:
                        f.write(page_text)
            
            logger.info(f"Successfully processed PDF with {result['total_pages']} pages")
            
            # Save the complete extraction result
            result_filename = f"{os.path.splitext(os.path.basename(file_path))[0]}_extraction.json"
            result_path = os.path.join(self.extracted_dir, result_filename)
            with open(result_path, 'w', encoding='utf-8') as f:
                json.dump(result, f, indent=2)
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing PDF: {str(e)}")
            result["error"] = str(e)
            return result
    
    def process_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text and metadata from a Word document.
        
        Args:
            file_path: Path to the Word document
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        logger.info(f"Processing Word document: {file_path}")
        
        result = {
            "file_name": os.path.basename(file_path),
            "file_path": file_path,
            "file_type": "docx",
            "paragraphs": [],
            "metadata": {},
            "tables": [],
            "processed_at": datetime.now().isoformat()
        }
        
        try:
            doc = docx.Document(file_path)
            
            # Extract metadata (core properties)
            core_props = doc.core_properties
            if core_props:
                result["metadata"] = {
                    "author": core_props.author or "",
                    "title": core_props.title or "",
                    "subject": core_props.subject or "",
                    "created": str(core_props.created) if core_props.created else "",
                    "modified": str(core_props.modified) if core_props.modified else "",
                    "category": core_props.category or ""
                }
            
            # Extract paragraphs
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    paragraph_data = {
                        "index": i,
                        "text": para.text,
                        "style": para.style.name if para.style else "Normal"
                    }
                    result["paragraphs"].append(paragraph_data)
            
            # Extract tables
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text)
                    table_data.append(row_data)
                
                table_info = {
                    "index": i,
                    "rows": len(table.rows),
                    "columns": len(table.rows[0].cells) if table.rows else 0,
                    "data": table_data
                }
                result["tables"].append(table_info)
            
            # Combine all text for full-text search
            all_text = "\n".join([para["text"] for para in result["paragraphs"]])
            
            # Save the extracted text to a file
            text_filename = f"{os.path.splitext(os.path.basename(file_path))[0]}.txt"
            text_path = os.path.join(self.extracted_dir, text_filename)
            with open(text_path, 'w', encoding='utf-8') as f:
                f.write(all_text)
            
            logger.info(f"Successfully processed Word document with {len(result['paragraphs'])} paragraphs and {len(result['tables'])} tables")
            
            # Save the complete extraction result
            result_filename = f"{os.path.splitext(os.path.basename(file_path))[0]}_extraction.json"
            result_path = os.path.join(self.extracted_dir, result_filename)
            with open(result_path, 'w', encoding='utf-8') as f:
                json.dump(result, f, indent=2)
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing Word document: {str(e)}")
            result["error"] = str(e)
            return result
    
    def process_image(self, file_path: str, detect_text: bool = False) -> Dict[str, Any]:
        """
        Process and analyze an image file.
        
        Args:
            file_path: Path to the image file
            detect_text: Whether to attempt OCR on the image (if available)
            
        Returns:
            Dictionary containing image information and extracted text if applicable
        """
        logger.info(f"Processing image: {file_path}")
        
        result = {
            "file_name": os.path.basename(file_path),
            "file_path": file_path,
            "file_type": "image",
            "image_info": {},
            "extracted_text": None,
            "processed_at": datetime.now().isoformat()
        }
        
        try:
            with Image.open(file_path) as img:
                # Get basic image information
                result["image_info"] = {
                    "format": img.format,
                    "mode": img.mode,
                    "size": img.size,
                    "width": img.width,
                    "height": img.height
                }
                
                # Get EXIF data if available
                if hasattr(img, '_getexif') and callable(img._getexif):
                    exif = img._getexif()
                    if exif:
                        exif_data = {}
                        for tag_id, value in exif.items():
                            try:
                                exif_data[str(tag_id)] = str(value)
                            except:
                                pass  # Skip problematic EXIF tags
                        result["image_info"]["exif"] = exif_data
                
                # Implement OCR if requested and available
                if detect_text:
                    try:
                        import pytesseract
                        result["extracted_text"] = pytesseract.image_to_string(img)
                        
                        # Save the extracted text
                        if result["extracted_text"]:
                            text_filename = f"{os.path.splitext(os.path.basename(file_path))[0]}_ocr.txt"
                            text_path = os.path.join(self.extracted_dir, text_filename)
                            with open(text_path, 'w', encoding='utf-8') as f:
                                f.write(result["extracted_text"])
                    except ImportError:
                        logger.warning("OCR requested but pytesseract is not installed")
                        result["extracted_text"] = None
                
                logger.info(f"Successfully processed image: {img.format} {img.size}")
            
            # Save the processing result
            result_filename = f"{os.path.splitext(os.path.basename(file_path))[0]}_info.json"
            result_path = os.path.join(self.extracted_dir, result_filename)
            with open(result_path, 'w', encoding='utf-8') as f:
                json.dump(result, f, indent=2)
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing image: {str(e)}")
            result["error"] = str(e)
            return result
    
    def process_base64_document(self, base64_data: str, 
                               file_name: str, 
                               file_type: str) -> Dict[str, Any]:
        """
        Process a document from base64-encoded data.
        
        Args:
            base64_data: Base64-encoded file data
            file_name: Original file name
            file_type: Type of file ("pdf", "docx", "image")
            
        Returns:
            Dictionary containing processing results
        """
        logger.info(f"Processing base64 document: {file_name}")
        
        try:
            # Decode base64 data
            file_data = base64.b64decode(base64_data)
            
            # Create a temporary file
            with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                temp_file.write(file_data)
                temp_path = temp_file.name
            
            # Process the file based on its type
            try:
                if file_type.lower() == "pdf":
                    result = self.process_pdf(temp_path)
                elif file_type.lower() == "docx":
                    result = self.process_docx(temp_path)
                elif file_type.lower() in ["image", "jpg", "jpeg", "png", "gif"]:
                    result = self.process_image(temp_path)
                else:
                    result = {
                        "error": f"Unsupported file type: {file_type}"
                    }
                
                # Add file name to result
                result["file_name"] = file_name
                
                # Copy the file to our documents directory for future reference
                dest_path = os.path.join(self.documents_dir, file_name)
                with open(dest_path, 'wb') as f:
                    f.write(file_data)
                
                result["saved_path"] = dest_path
                
            finally:
                # Clean up the temporary file
                try:
                    os.unlink(temp_path)
                except:
                    pass
            
            return result
            
        except Exception as e:
            logger.error(f"Error processing base64 document: {str(e)}")
            return {"error": str(e)}
    
    def extract_tables_from_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extract all tables from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            List of dictionaries containing tables by page
        """
        logger.info(f"Extracting tables from PDF: {file_path}")
        
        results = []
        
        try:
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    try:
                        tables = page.extract_tables()
                        if tables:
                            table_data = []
                            for j, table in enumerate(tables):
                                if table:
                                    clean_table = []
                                    for row in table:
                                        if row:
                                            clean_table.append([cell or "" for cell in row])
                                    
                                    table_data.append({
                                        "table_index": j,
                                        "data": clean_table
                                    })
                            
                            if table_data:
                                results.append({
                                    "page_number": i + 1,
                                    "tables": table_data
                                })
                    except Exception as e:
                        logger.warning(f"Error extracting tables from page {i+1}: {str(e)}")
            
            # Save the extraction result
            result_filename = f"{os.path.splitext(os.path.basename(file_path))[0]}_tables.json"
            result_path = os.path.join(self.extracted_dir, result_filename)
            with open(result_path, 'w', encoding='utf-8') as f:
                json.dump(results, f, indent=2)
            
            logger.info(f"Extracted tables from PDF: {len(results)} pages with tables")
            
            return results
            
        except Exception as e:
            logger.error(f"Error extracting tables from PDF: {str(e)}")
            return []
    
    def merge_pdfs(self, pdf_paths: List[str], output_path: Optional[str] = None) -> str:
        """
        Merge multiple PDFs into a single PDF.
        
        Args:
            pdf_paths: List of paths to PDF files
            output_path: Path for the output file (optional)
            
        Returns:
            Path to the merged PDF file
        """
        logger.info(f"Merging {len(pdf_paths)} PDFs")
        
        try:
            # We'll use PyPDF2 if available, otherwise fallback to a shell command
            try:
                from PyPDF2 import PdfMerger
                
                merger = PdfMerger()
                for pdf_path in pdf_paths:
                    merger.append(pdf_path)
                
                if not output_path:
                    output_path = os.path.join(self.output_dir, f"merged_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf")
                
                merger.write(output_path)
                merger.close()
                
                logger.info(f"Successfully merged PDFs to: {output_path}")
                return output_path
                
            except ImportError:
                # Fallback to pdftk if installed
                import subprocess
                
                if not output_path:
                    output_path = os.path.join(self.output_dir, f"merged_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf")
                
                # Check if pdftk is installed
                try:
                    subprocess.run(["pdftk", "--version"], capture_output=True, check=True)
                    
                    # Construct pdftk command
                    cmd = ["pdftk"]
                    cmd.extend(pdf_paths)
                    cmd.extend(["cat", "output", output_path])
                    
                    # Run the command
                    subprocess.run(cmd, check=True)
                    
                    logger.info(f"Successfully merged PDFs using pdftk to: {output_path}")
                    return output_path
                    
                except (subprocess.SubprocessError, FileNotFoundError):
                    raise ImportError("Neither PyPDF2 nor pdftk is available for merging PDFs")
                
        except Exception as e:
            logger.error(f"Error merging PDFs: {str(e)}")
            return f"Error: {str(e)}"
    
    def split_pdf(self, pdf_path: str, pages: List[int], output_path: Optional[str] = None) -> str:
        """
        Extract specific pages from a PDF into a new PDF.
        
        Args:
            pdf_path: Path to the PDF file
            pages: List of page numbers to extract (1-based index)
            output_path: Path for the output file (optional)
            
        Returns:
            Path to the extracted PDF file
        """
        logger.info(f"Splitting PDF {pdf_path}, extracting pages {pages}")
        
        try:
            # We'll use PyPDF2 if available
            try:
                from PyPDF2 import PdfReader, PdfWriter
                
                reader = PdfReader(pdf_path)
                writer = PdfWriter()
                
                for page_num in pages:
                    # Convert from 1-based to 0-based indexing
                    idx = page_num - 1
                    if 0 <= idx < len(reader.pages):
                        writer.add_page(reader.pages[idx])
                    else:
                        logger.warning(f"Page {page_num} is out of range")
                
                if not output_path:
                    pages_str = "-".join(str(p) for p in pages[:5])
                    if len(pages) > 5:
                        pages_str += f"-and-{len(pages)-5}-more"
                    
                    output_path = os.path.join(
                        self.output_dir, 
                        f"{os.path.splitext(os.path.basename(pdf_path))[0]}_pages_{pages_str}.pdf"
                    )
                
                with open(output_path, "wb") as output_file:
                    writer.write(output_file)
                
                logger.info(f"Successfully extracted {len(pages)} pages to: {output_path}")
                return output_path
                
            except ImportError:
                # If PyPDF2 is not available, we'll use pdftk if installed
                import subprocess
                
                if not output_path:
                    pages_str = "-".join(str(p) for p in pages[:5])
                    if len(pages) > 5:
                        pages_str += f"-and-{len(pages)-5}-more"
                    
                    output_path = os.path.join(
                        self.output_dir, 
                        f"{os.path.splitext(os.path.basename(pdf_path))[0]}_pages_{pages_str}.pdf"
                    )
                
                # Check if pdftk is installed
                try:
                    subprocess.run(["pdftk", "--version"], capture_output=True, check=True)
                    
                    # Construct pdftk command
                    pages_str = " ".join(str(p) for p in pages)
                    cmd = ["pdftk", pdf_path, "cat", pages_str, "output", output_path]
                    
                    # Run the command
                    subprocess.run(cmd, check=True)
                    
                    logger.info(f"Successfully extracted pages using pdftk to: {output_path}")
                    return output_path
                    
                except (subprocess.SubprocessError, FileNotFoundError):
                    raise ImportError("Neither PyPDF2 nor pdftk is available for splitting PDFs")
                
        except Exception as e:
            logger.error(f"Error splitting PDF: {str(e)}")
            return f"Error: {str(e)}"
    
    def crop_image(self, image_path: str, crop_box: Tuple[int, int, int, int], 
                  output_path: Optional[str] = None) -> str:
        """
        Crop an image to the specified bounding box.
        
        Args:
            image_path: Path to the image file
            crop_box: Tuple of (left, top, right, bottom) coordinates
            output_path: Path for the output file (optional)
            
        Returns:
            Path to the cropped image file
        """
        logger.info(f"Cropping image {image_path} with box {crop_box}")
        
        try:
            with Image.open(image_path) as img:
                # Crop the image
                cropped_img = img.crop(crop_box)
                
                # Determine output path
                if not output_path:
                    output_path = os.path.join(
                        self.output_dir, 
                        f"{os.path.splitext(os.path.basename(image_path))[0]}_cropped.{img.format.lower()}"
                    )
                
                # Save the cropped image
                cropped_img.save(output_path)
                
                logger.info(f"Successfully cropped image to: {output_path}")
                return output_path
                
        except Exception as e:
            logger.error(f"Error cropping image: {str(e)}")
            return f"Error: {str(e)}"
    
    def resize_image(self, image_path: str, new_size: Tuple[int, int], 
                    output_path: Optional[str] = None) -> str:
        """
        Resize an image to the specified dimensions.
        
        Args:
            image_path: Path to the image file
            new_size: Tuple of (width, height) in pixels
            output_path: Path for the output file (optional)
            
        Returns:
            Path to the resized image file
        """
        logger.info(f"Resizing image {image_path} to {new_size}")
        
        try:
            with Image.open(image_path) as img:
                # Resize the image
                resized_img = img.resize(new_size, Image.Resampling.LANCZOS)
                
                # Determine output path
                if not output_path:
                    width, height = new_size
                    output_path = os.path.join(
                        self.output_dir, 
                        f"{os.path.splitext(os.path.basename(image_path))[0]}_{width}x{height}.{img.format.lower()}"
                    )
                
                # Save the resized image
                resized_img.save(output_path)
                
                logger.info(f"Successfully resized image to: {output_path}")
                return output_path
                
        except Exception as e:
            logger.error(f"Error resizing image: {str(e)}")
            return f"Error: {str(e)}"

# For testing
if __name__ == "__main__":
    processor = DocumentProcessor()
    
    # Test with a sample PDF file if one exists
    pdf_files = [f for f in os.listdir(".") if f.endswith(".pdf")]
    if pdf_files:
        sample_pdf = pdf_files[0]
        processor.process_pdf(sample_pdf)
        
    # Test with a sample DOCX file if one exists
    docx_files = [f for f in os.listdir(".") if f.endswith(".docx")]
    if docx_files:
        sample_docx = docx_files[0]
        processor.process_docx(sample_docx)
        
    # Test with a sample image file if one exists
    image_files = [f for f in os.listdir(".") if f.lower().endswith((".jpg", ".jpeg", ".png"))]
    if image_files:
        sample_image = image_files[0]
        processor.process_image(sample_image)